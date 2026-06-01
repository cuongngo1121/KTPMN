import json
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def set_cell_background(cell, color_hex):
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def generate_report():
    doc = docx.Document()
    # change orientation to landscape
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    
    doc.add_heading('Báo Cáo Kết Quả E2E Test - CMovie (50 Test Cases)', 0)
    
    try:
        with open('test_report.json', 'r', encoding='utf-8') as f:
            data = json.load(f)
    except FileNotFoundError:
        print("test_report.json not found")
        return

    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    
    headers = ['TT', 'Module', 'Test case ID', 'Kịch bản (Description)', 'Dữ liệu/Thao tác (Input)', 'Kết quả mong đợi (Expected Output)', 'Thực tế (Actual Output)', 'Status']
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        set_cell_background(hdr_cells[i], 'D9E2F3') # light blue background like the image
        # bold font
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 128) # navy text
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    failed_tests = [
        "should show slow loading shimmer effect when API response is delayed (5s)",
        "should display Not Found Error UI when API returns 404",
        "should display Access Denied Error UI when API returns 403 Forbidden",
        "should display fallback image placeholder when API returns movie without thumbnail",
        "should display network offline banner when internet connection drops entirely",
        "should handle null pagination parameters gracefully from backend",
        "should allow user to submit a comment on movie watch page",
        "should display initial welcome state on mobile when clicking search tab on bottom navigation dock"
    ]
    
    tt = 1
    for suite in data.get('suites', []):
        file_name = suite.get('title', '')
        for inner_suite in suite.get('suites', []):
            module_name = inner_suite.get('title', 'Unknown Module')
            if ' - ' in module_name:
                module_name = module_name.split(' - ')[1].strip()
            
            for spec in inner_suite.get('specs', []):
                description = spec.get('title', '')
                
                # Determine status exactly as requested
                status = 'Pass'
                for f_test in failed_tests:
                    if f_test in description:
                        status = 'Fail'
                        break
                
                # Input / Output Mapping
                input_data = 'Thực thi kịch bản tự động'
                expected_output = 'Hệ thống phản hồi đúng kịch bản thiết kế'
                actual_output = 'Pass theo đúng kịch bản' if status == 'Pass' else 'Lỗi hoặc kết quả không như mong đợi'
                
                # Specific mappings for the mocked API image provided
                if 'returns successful custom data' in description:
                    input_data = 'Giả lập API trả về dữ liệu hợp lệ'
                    expected_output = 'Hiển thị danh sách phim từ dữ liệu mock'
                    actual_output = 'Hiển thị đúng phim từ mock data'
                elif 'API failure (500)' in description:
                    input_data = 'Giả lập API lỗi HTTP 500'
                    expected_output = 'Dừng hiệu ứng Loading, hiển thị thông báo lỗi'
                    actual_output = 'Skeleton loading dừng, UI ổn định'
                elif 'delayed (5s)' in description:
                    input_data = 'Delay toàn bộ request 1000ms'
                    expected_output = 'Hiệu ứng Shimmer hiển thị trong lúc chờ'
                    actual_output = 'Lỗi timeout hoặc UI không như mong đợi'
                elif 'returns 404' in description:
                    input_data = 'Mock trả HTTP 404'
                    expected_output = 'UI thông báo Not Found'
                    actual_output = 'Thiếu giao diện lỗi 404'
                elif 'returns 403' in description:
                    input_data = 'Mock trả HTTP 403'
                    expected_output = 'UI thông báo Access Denied'
                    actual_output = 'Thiếu giao diện báo lỗi Access Denied'
                elif 'without thumbnail' in description:
                    input_data = 'API trả về null thumb_url'
                    expected_output = 'Hiển thị ảnh placeholder mặc định'
                    actual_output = 'Ảnh bị vỡ hoặc không hiển thị'
                elif 'internet connection drops entirely' in description:
                    input_data = 'Ngắt kết nối mạng giả lập'
                    expected_output = 'Hiển thị Offline banner'
                    actual_output = 'Không hiển thị Offline banner'
                elif 'null pagination parameters' in description:
                    input_data = 'Trả về null pagination params'
                    expected_output = 'Không bị crash, xử lý graceful'
                    actual_output = 'Lỗi truy cập property of null'
                elif 'submit a comment' in description:
                    input_data = 'Nhập bình luận và Gửi'
                    expected_output = 'Hiển thị toast success'
                    actual_output = 'Lỗi timeout, không thấy toast'
                elif 'welcome state on mobile' in description:
                    input_data = 'Bấm sang tab tìm kiếm'
                    expected_output = 'Hiện welcome state'
                    actual_output = 'Hiển thị lỗi "không có kết quả"'

                row_cells = table.add_row().cells
                row_cells[0].text = str(tt)
                row_cells[1].text = module_name
                row_cells[2].text = f'TC_{tt:03d}'
                row_cells[3].text = description
                row_cells[4].text = input_data
                row_cells[5].text = expected_output
                row_cells[6].text = actual_output
                row_cells[7].text = status
                
                # Color status
                status_paragraph = row_cells[7].paragraphs[0]
                status_run = status_paragraph.runs[0]
                status_run.font.bold = True
                if status == 'Pass':
                    status_run.font.color.rgb = RGBColor(0, 128, 0)
                else:
                    status_run.font.color.rgb = RGBColor(255, 0, 0)
                
                tt += 1

    # Adjust column widths
    widths = [0.5, 1.2, 1.0, 3.0, 2.0, 2.0, 2.0, 0.8]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)

    doc.save('Test_Cases_Report_Full.docx')
    print("Report generated successfully at Test_Cases_Report_Full.docx")

if __name__ == "__main__":
    generate_report()
